import os
import re

# We can try to import openpyxl or docx if they are installed, otherwise use raw text/regex if they are zip files, or use python-docx / openpyxl if available.
# Let's check if they are installed by trying to import them. If not, we can read the raw zip contents since docx and xlsx are zip archives containing XML files.
try:
    import openpyxl
    has_openpyxl = True
except ImportError:
    has_openpyxl = False

try:
    import docx
    has_docx = True
except ImportError:
    has_docx = False

import zipfile
import xml.etree.ElementTree as ET

def search_xlsx(path, keyword_pat):
    if has_openpyxl:
        try:
            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                for r, row in enumerate(ws.iter_rows(values_only=True)):
                    for c, val in enumerate(row):
                        if val and keyword_pat.search(str(val)):
                            print(f"[XLSX MATCH] {path} | Sheet: {sheet} | Row {r+1}, Col {c+1}: {val}")
            return
        except Exception as e:
            pass # fall back to zip parsing
    
    # fallback to raw xml parsing
    try:
        with zipfile.ZipFile(path) as z:
            # Shared strings
            shared_strings = []
            if 'xl/sharedStrings.xml' in z.namelist():
                ss_content = z.read('xl/sharedStrings.xml')
                root = ET.fromstring(ss_content)
                for t in root.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}t'):
                    shared_strings.append(t.text or '')
            
            # Check sheet data
            for name in z.namelist():
                if name.startswith('xl/worksheets/sheet'):
                    sheet_content = z.read(name)
                    root = ET.fromstring(sheet_content)
                    for cell in root.findall('.//{http://schemas.openxmlformats.org/spreadsheetml/2006/main}v'):
                        val = cell.text or ''
                        # if it's a shared string index
                        t = cell.get('t')
                        if t == 's':
                            try:
                                idx = int(val)
                                if 0 <= idx < len(shared_strings):
                                    val = shared_strings[idx]
                            except ValueError:
                                pass
                        if keyword_pat.search(val):
                            print(f"[XLSX MATCH (raw)] {path} | {name}: {val}")
    except Exception as e:
        print(f"Error reading xlsx {path}: {e}")

def search_docx(path, keyword_pat):
    if has_docx:
        try:
            doc = docx.Document(path)
            for p_idx, para in enumerate(doc.paragraphs):
                if keyword_pat.search(para.text):
                    print(f"[DOCX MATCH] {path} | Para {p_idx+1}: {para.text}")
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        if keyword_pat.search(cell.text):
                            print(f"[DOCX MATCH Table] {path} | Table {t_idx+1}, R{r_idx+1}C{c_idx+1}: {cell.text}")
            return
        except Exception as e:
            pass # fall back to zip parsing
            
    # fallback to raw xml
    try:
        with zipfile.ZipFile(path) as z:
            xml_content = z.read('word/document.xml')
            root = ET.fromstring(xml_content)
            # Find all text elements
            texts = []
            for t in root.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if t.text and keyword_pat.search(t.text):
                    print(f"[DOCX MATCH (raw)] {path}: {t.text}")
    except Exception as e:
        print(f"Error reading docx {path}: {e}")

search_dirs = [
    "/Users/haak78/Desktop",
    "/Users/haak78/Documents",
    "/Users/haak78/Library/CloudStorage/GoogleDrive-tomas.hajek.photographer@gmail.com/Můj disk"
]

patterns = [
    re.compile(r"pardubic", re.IGNORECASE),
    re.compile(r"martin", re.IGNORECASE),
    re.compile(r"newsletter", re.IGNORECASE)
]

print("Starting deep search in office files...")
for sdir in search_dirs:
    if not os.path.exists(sdir):
        continue
    for root, dirs, files in os.walk(sdir):
        for name in files:
            path = os.path.join(root, name)
            # check extensions
            ext = os.path.splitext(name)[1].lower()
            for pat in patterns:
                if ext in ['.xlsx', '.xlsm']:
                    search_xlsx(path, pat)
                elif ext in ['.docx']:
                    search_docx(path, pat)
                elif ext in ['.txt', '.csv']:
                    try:
                        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                            for l_idx, line in enumerate(f):
                                if pat.search(line):
                                    print(f"[TXT MATCH] {path} | Line {l_idx+1}: {line.strip()}")
                    except Exception as e:
                        pass
print("Done deep search.")
